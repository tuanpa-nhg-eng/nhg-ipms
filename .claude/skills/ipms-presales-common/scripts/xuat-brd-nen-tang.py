#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Sinh BRD NỀN TẢNG iPMS (AI-Native Platform) — tài liệu NỘI BỘ để xây nền tảng.

    py xuat-brd-nen-tang.py <brd-nen-tang.json> [--out <thư-mục>]

Đầu ra:
    01_BRD_iPMS_AI_Native_{v}.docx    — 18 mục
    02_Ma_tran_Yeu_cau_iPMS_{v}.xlsx  — 13 sheet

KHÁC với xuat-ho-so.py (BRD tiền bán hàng gửi khách ngoài):
  · Đây là tài liệu nội bộ → KHÔNG áp hàng rào "không nhắc NHG/khối B0-B5/trạng thái xây dựng".
  · Vẫn loại khoá "_" (ghi chú soạn thảo) để tài liệu sạch.
Trục xương sống: Mục 4 — Từ điển Tác vụ gắn KPI là LÕI nền tảng.
"""
import argparse, json, sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

ROOT = Path(__file__).resolve().parents[4]
LOGO = ROOT / "design-system" / "public" / "logo.png"

GREEN, GREEN_D, RED, INK, INK_3 = "037236", "04361C", "ED2024", "101214", "7C8285"
GREEN_50, MIST, WARN_50, INFO_50, RED_50, GOLD_50 = "E8F4ED", "EEF3EF", "FBF0DC", "E5F0F9", "FDECEC", "FFF8E1"
FONT = "Be Vietnam Pro"

TT_MAU = {"Đã có": GREEN_50, "Một phần": WARN_50, "Chưa có": RED_50}
UU_TEN = {"M": "Bắt buộc", "S": "Nên có", "C": "Có thì tốt", "W": "Đợt sau"}

# Chế độ BẢN TRỐNG — đặt trong main(). Khi bật, mọi ô rỗng biến thành chỗ để điền
# thay vì bị bỏ qua, và mỗi mục có thêm dòng hướng dẫn cần điền gì.
TRONG = False
SO_DONG_TRONG = 6


def bo_khoa_soan_thao(x):
    if isinstance(x, dict):
        return {k: bo_khoa_soan_thao(v) for k, v in x.items() if not k.startswith("_")}
    if isinstance(x, list):
        return [bo_khoa_soan_thao(v) for v in x]
    return x


# Khoá GIỮ LẠI khi làm trống — đây là phần khung/câu hỏi gợi ý, không phải câu trả lời.
GIU = {
    "muoi_hai_cau_hoi": ["stt", "cau_hoi"],
    "loi_nen_tang.chuoi_truy_vet": ["mat_xich", "cau_hoi"],
    "ranh_gioi_ai.cong_hitl": ["cong"],
    "phi_chuc_nang": ["nhom"],
    "du_lieu": ["nhom"],
    "lo_trinh": ["giai_doan"],
    "trung_tam_quyen_luc": ["ma", "ten", "vai_tro"],
    "modules": ["ma", "ten"],
    "persona": ["persona"],
    "boi_canh.gia_tri_3v": ["gia_tri"],
    "tai_lieu": ["ten", "ten_en"],
}


def lam_trong(x, duong=""):
    """Xoá mọi câu trả lời, giữ nguyên khung: tiêu đề mục, câu hỏi gợi ý, tên mắt xích."""
    if isinstance(x, dict):
        giu = GIU.get(duong, [])
        return {k: (v if k in giu else lam_trong(v, f"{duong}.{k}" if duong else k))
                for k, v in x.items()}
    if isinstance(x, list):
        giu = GIU.get(duong)
        if giu:                                    # danh sách khung — giữ dòng, xoá câu trả lời
            return [lam_trong(v, duong) for v in x]
        return []                                  # danh sách nội dung — xoá sạch, sinh dòng trống
    return ""


# Hướng dẫn điền theo từng mục — chỉ hiện ở bản trống.
HD = {
    1: "Nêu tài liệu này phục vụ quyết định gì, ai đọc, và ranh giới với tài liệu thiết kế kỹ thuật.",
    2: "Trả lời: tổ chức đang ở đâu, vấn đề gì đang tồn tại, và hệ giá trị được chuyển thành nguyên tắc thiết kế ra sao. Tránh viết chung chung — mỗi vấn đề phải có dấu hiệu quan sát được.",
    3: "Định vị bằng một câu người ngoài đọc cũng hiểu. Mỗi mục tiêu phải có CÁCH ĐO, người chịu trách nhiệm và mốc thời gian — mục tiêu không đo được thì không đưa vào.",
    4: "★ Mục quan trọng nhất. Xác định LÕI của nền tảng: thứ mà mất nó thì sản phẩm mất lý do tồn tại. Viết rõ vì sao là lõi, chuỗi truy vết đi qua những mắt xích nào, và ràng buộc nào bất biến.",
    5: "Ghi cả hai chiều. Phần NGOÀI phạm vi quan trọng ngang phần trong — mỗi hạng mục ngoài phạm vi phải nói rõ giao diện với hệ nào.",
    6: "Mười hai câu hỏi bắt buộc theo NHG Strategic Context §10. Không được để trống câu nào; câu chưa trả lời được thì ghi rõ ai sẽ trả lời và khi nào.",
    7: "Với mỗi khối chức năng: quan hệ với hệ thống là gì và cần dashboard nào. Với mỗi nhóm người dùng: làm được gì và đặc biệt là KHÔNG được làm gì.",
    8: "Chia hệ thống thành module theo năng lực nghiệp vụ, không theo màn hình. Đánh dấu module nào là lõi.",
    9: "Mỗi yêu cầu: một câu bắt đầu bằng “Hệ thống phải…”, kèm tiêu chí chấp nhận KIỂM CHỨNG ĐƯỢC khi nghiệm thu. Tiêu chí kiểu “hoạt động tốt” là tiêu chí rỗng.",
    10: "Liệt kê tường minh AI được phép và không được phép làm gì, rồi xác định các cổng bắt buộc có người quyết định. Danh sách này phải cưỡng chế được bằng mã, không phải bằng hướng dẫn sử dụng.",
    11: "Mỗi agent phải khai đủ: mục đích, dữ liệu dùng, mức phân loại tối đa được xử lý, cổng người duyệt và chủ quản. Thiếu trường nào thì agent đó chưa được phép chạy.",
    12: "Với mỗi nhóm dữ liệu: nguồn gốc, chủ dữ liệu, mức phân loại và ranh giới AI. Đây là căn cứ để cấu hình chính sách xuất dữ liệu.",
    13: "Yêu cầu phi chức năng phải có mục tiêu đo được, không viết “nhanh”, “ổn định”.",
    14: "Mỗi giai đoạn phải có CỔNG RA kiểm chứng được bởi người ngoài đội xây dựng — không phải tự đánh giá.",
    15: "Mỗi rủi ro: hiện tượng → mức → biện pháp giảm thiểu → người chịu trách nhiệm. Rủi ro không có người chịu trách nhiệm là rủi ro không được quản lý.",
    16: "Giả định là điều đang coi là đúng mà chưa kiểm chứng. Luôn ghi ảnh hưởng nếu giả định sai.",
    17: "Điểm còn chờ quyết định, ai quyết, và yêu cầu nào đang bị chặn vì nó.",
    18: "Xác định đủ ba tuyến phê duyệt: chủ quản nghiệp vụ, chủ quản công nghệ và cấp phê duyệt cuối.",
}


# ─────────────────────────── DOCX ───────────────────────────
def _field(par, ma):
    r = par.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    t = OxmlElement("w:instrText"); t.set(qn("xml:space"), "preserve"); t.text = ma
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    r._r.append(b); r._r.append(t); r._r.append(e)


def _to_cell(cell, mau):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), mau)
    tcPr.append(shd)


def _txt(par, text, size=10.5, bold=False, italic=False, mau=INK):
    r = par.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = RGBColor.from_string(mau)
    return r


def _o_viet(doc, dong=3):
    """Ô trống để viết vào — chỉ dùng ở bản trống."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; c.text = ""
    _to_cell(c, "FCFCFA")
    for _ in range(dong - 1):
        c.add_paragraph()
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _muc(doc, so, vi, en="", ngat=False):
    if ngat:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    _txt(p, f"MỤC {so}. ", 8.5, True, mau=RED)
    _txt(p, vi, 15, True, mau=GREEN_D)
    if en:
        q = doc.add_paragraph(); q.paragraph_format.space_after = Pt(7)
        q.paragraph_format.keep_with_next = True
        _txt(q, en, 9, italic=True, mau=INK_3)
    if TRONG and so in HD:
        h = doc.add_paragraph(); h.paragraph_format.space_after = Pt(8)
        h.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _txt(h, "▸ Cần điền: ", 9, True, mau=RED)
        _txt(h, HD[so], 9, italic=True, mau=INK_3)


def _sub(doc, vi):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    _txt(p, vi, 11.5, True, mau=GREEN)


def _para(doc, text, size=10.5, italic=False, mau=INK, before=0, after=6, bold=False):
    if not text:                       # ô dữ liệu rỗng
        if TRONG:
            _o_viet(doc, 3)
        return None
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _txt(p, text, size, bold=bold, italic=italic, mau=mau)
    return p


def _cham(doc, items, size=10):
    if not items and TRONG:
        items = ["" for _ in range(4)]
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _txt(p, it if it else " " * 90, size, mau=INK if it else INK_3)


def _bang(doc, dau, dong, rong=None, mau_o=None):
    if not dong and TRONG:             # bảng chưa có dữ liệu → sinh dòng trống để điền
        dong = [["" for _ in dau] for _ in range(SO_DONG_TRONG)]
    t = doc.add_table(rows=1, cols=len(dau)); t.style = "Table Grid"
    for i, h in enumerate(dau):
        c = t.rows[0].cells[i]; c.text = ""
        _txt(c.paragraphs[0], h, 9, True, mau="FFFFFF")
        _to_cell(c, GREEN)
    for ri, r in enumerate(dong):
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = ""
            _txt(cells[i].paragraphs[0], "" if v is None else str(v), 9)
            if mau_o:
                m = mau_o(ri, i, v)
                if m:
                    _to_cell(cells[i], m)
    if rong:
        for i, w in enumerate(rong):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _khung(doc, tieu_de, noi_dung, mau_nen=GOLD_50):
    """Khối nhấn mạnh — dùng cho tuyên bố lõi nền tảng."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    c = t.rows[0].cells[0]; c.text = ""
    _to_cell(c, mau_nen)
    p = c.paragraphs[0]
    _txt(p, tieu_de + "\n", 11.5, True, mau=GREEN_D)
    if noi_dung:
        _txt(p, noi_dung, 10.5)
    elif TRONG:
        _txt(p, "……………………………………………………………………………………………………", 10.5, mau=INK_3)
        c.add_paragraph(); c.add_paragraph()
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def dung_docx(d, ra: Path):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT; st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(1.8)

    tl = d["tai_lieu"]
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _txt(fp, f"{tl.get('ma','')} · {tl.get('phan_loai','Nội bộ')} · Trang ", 8, mau=INK_3)
    _field(fp, "PAGE"); _txt(fp, " / ", 8, mau=INK_3); _field(fp, "NUMPAGES")
    for r in fp.runs:
        r.font.size = Pt(8); r.font.name = FONT; r.font.color.rgb = RGBColor.from_string(INK_3)

    # ══ BÌA ══
    if LOGO.exists():
        pl = doc.add_paragraph(); pl.add_run().add_picture(str(LOGO), height=Cm(1.25))
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(50)
    _txt(p, "TÀI LIỆU YÊU CẦU NGHIỆP VỤ", 24, True, mau=GREEN_D)
    p = doc.add_paragraph(); _txt(p, "NỀN TẢNG iPMS — AI-NATIVE PLATFORM", 15, True, mau=RED)
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6)
    _txt(p, "Business Requirements Document — iPMS AI-Native Platform", 10.5, italic=True, mau=INK_3)
    if TRONG:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14)
        _txt(p, "BẢN TRỐNG — ĐỂ ĐIỀN", 14, True, mau=RED)
        p = doc.add_paragraph()
        _txt(p, "Khuôn tài liệu giữ nguyên cấu trúc 18 mục. Mỗi mục có dòng “▸ Cần điền” nêu rõ cần viết gì. "
                "Ô nền nhạt và dòng bảng trống là chỗ điền. Xoá dòng hướng dẫn sau khi điền xong.",
             9.5, italic=True, mau=INK_3)
    else:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(20)
        _txt(p, "Hệ điều hành hiệu suất và tăng trưởng của Tập đoàn Nguyễn Hoàng", 12.5, mau=INK)
        p = doc.add_paragraph()
        _txt(p, "Chiến lược đi xuống thành tác vụ · tác vụ tạo ra dữ liệu · dữ liệu nuôi dashboard · "
                "dashboard hỗ trợ quyết định · AI tăng tốc con người · governance bảo vệ niềm tin", 9.5,
             italic=True, mau=GREEN)
    doc.add_paragraph().paragraph_format.space_before = Pt(16)
    _bang(doc, ["", ""], [
        ["Mã tài liệu", tl.get("ma", "")],
        ["Phiên bản", f"{tl.get('phien_ban','')} · {tl.get('ngay','')}"],
        ["Chủ quản nghiệp vụ", tl.get("chu_quan_nghiep_vu", "")],
        ["Chủ quản công nghệ", tl.get("chu_quan_cong_nghe", "")],
        ["Cấp phê duyệt", tl.get("cap_phe_duyet", "")],
        ["Phân loại", tl.get("phan_loai", "")],
        ["Số yêu cầu", f"{len(d.get('yeu_cau', []))} yêu cầu / {len(d.get('modules', []))} module"],
    ], rong=[5.0, 11.5])
    _sub(doc, "Căn cứ lập tài liệu")
    _cham(doc, tl.get("can_cu", []), size=9.5)

    # ══ 1 ══
    _muc(doc, 1, "Mục đích & phạm vi tài liệu", "Purpose and document scope", ngat=True)
    _para(doc, "Tài liệu này xác lập yêu cầu nghiệp vụ để xây dựng nền tảng iPMS như một AI-Native Platform "
               "của Tập đoàn. Tài liệu trả lời câu hỏi CÁI GÌ và VÌ SAO; thiết kế kỹ thuật, ước lượng nguồn lực "
               "và kế hoạch triển khai nằm ở các tài liệu riêng.")
    _para(doc, "Đối tượng đọc: Ban Điều hành, các khối chức năng V1, B0, B1, B2, B3, B5, B6 và đơn vị vận hành. "
               "Tài liệu là căn cứ để thống nhất phạm vi trước khi phân bổ nguồn lực.")

    # ══ 2 ══
    bc = d.get("boi_canh", {})
    _muc(doc, 2, "Bối cảnh chiến lược", "Strategic context")
    _sub(doc, "Cách hiểu về Tập đoàn")
    _para(doc, bc.get("dinh_vi_nhg", ""))
    _sub(doc, "Vấn đề đang tồn tại")
    _para(doc, bc.get("van_de", ""))
    _sub(doc, "Hệ giá trị 3V chuyển thành nguyên tắc thiết kế")
    _bang(doc, ["Giá trị", "Diễn giải trong iPMS"],
          [[x.get("gia_tri", ""), x.get("trong_ipms", "")] for x in bc.get("gia_tri_3v", [])],
          rong=[4.2, 12.3])

    # ══ 3 ══
    _muc(doc, 3, "Định vị sản phẩm & mục tiêu", "Product positioning and objectives")
    _khung(doc, "Định vị", bc.get("dinh_vi_ipms", ""), GREEN_50)
    _sub(doc, "Nguyên tắc kiến trúc bắt buộc")
    _cham(doc, bc.get("nguyen_tac_thiet_ke", []), size=9.8)
    _sub(doc, "Mục tiêu và cách đo")
    _bang(doc, ["Mã", "Mục tiêu", "Cách đo", "Chủ quản", "Mốc"],
          [[x.get("ma", ""), x.get("muc_tieu", ""), x.get("do_luong", ""), x.get("chu_quan", ""), x.get("moc", "")]
           for x in d.get("muc_tieu", [])], rong=[1.6, 4.4, 6.4, 1.9, 2.2])

    # ══ 4 — LÕI ══
    ln = d.get("loi_nen_tang", {})
    _muc(doc, 4, "LÕI NỀN TẢNG — " + (ln.get("tieu_de") or "……… (điền tên lõi) ………"),
         "Platform core", ngat=True)
    _khung(doc, "★ Tuyên bố lõi", ln.get("tuyen_bo", ""))
    _sub(doc, "Vì sao đây là lõi")
    _cham(doc, ln.get("luan_diem", []), size=10)
    _sub(doc, "Chuỗi truy vết — từ chiến lược tới việc làm và ngược lại")
    _bang(doc, ["Mắt xích", "Câu hỏi nó trả lời", "Chủ thể", "Ví dụ"],
          [[x.get("mat_xich", ""), x.get("cau_hoi", ""), x.get("chu_the", ""), x.get("vi_du", "")]
           for x in ln.get("chuoi_truy_vet", [])], rong=[3.0, 6.0, 3.4, 4.1],
          mau_o=lambda ri, ci, v: GREEN_50 if ci == 0 else None)
    _sub(doc, "Ràng buộc bất biến của lõi")
    _cham(doc, ln.get("rang_buoc_bat_bien", []), size=10)
    _sub(doc, "Vòng lặp tối ưu liên tục")
    _para(doc, ln.get("vong_lap_toi_uu", ""))
    _sub(doc, "Lõi này là nền để triển khai AI")
    _bang(doc, ["Năng lực", "Ý nghĩa"],
          [[x.get("nang_luc", ""), x.get("y_nghia", "")] for x in ln.get("vai_tro_ai", [])],
          rong=[6.0, 10.5])
    _khung(doc, "Hậu quả nếu bỏ lõi", ln.get("hau_qua_neu_bo_lo", ""), RED_50)

    # ══ 5 ══
    pv = d.get("pham_vi", {})
    _muc(doc, 5, "Phạm vi hệ thống", "System scope", ngat=True)
    _sub(doc, "Trong phạm vi")
    _cham(doc, pv.get("trong", []), size=9.8)
    _sub(doc, "Ngoài phạm vi và giao diện với hệ khác")
    _bang(doc, ["Hạng mục", "Lý do nằm ngoài", "Giao diện với iPMS"],
          [[x.get("hang_muc", ""), x.get("ly_do", ""), x.get("giao_dien", "")] for x in pv.get("ngoai", [])],
          rong=[4.4, 6.4, 5.7])

    # ══ 6 ══
    _muc(doc, 6, "Mười hai câu hỏi bắt buộc", "The twelve mandatory questions")
    _para(doc, "Mọi nền tảng của Tập đoàn phải trả lời được mười hai câu hỏi này trước khi được phê duyệt xây dựng. "
               "Đây là phần soi chiếu bắt buộc, không phải phần bổ sung.", 9.8, italic=True, mau=INK_3)
    _bang(doc, ["#", "Câu hỏi", "Trả lời của iPMS"],
          [[str(x.get("stt", "")), x.get("cau_hoi", ""), x.get("tra_loi", "")] for x in d.get("muoi_hai_cau_hoi", [])],
          rong=[1.0, 4.6, 10.9])

    # ══ 7 ══
    _muc(doc, 7, "Trung tâm quyền lực & người dùng", "Power centres and users", ngat=True)
    _sub(doc, "Quan hệ với các khối chức năng")
    _bang(doc, ["Mã", "Khối", "Quan hệ với iPMS", "Dashboard cần có"],
          [[x.get("ma", ""), x.get("ten", ""), x.get("quan_he_ipms", ""), x.get("dashboard", "")]
           for x in d.get("trung_tam_quyen_luc", [])], rong=[1.3, 3.8, 5.7, 5.7])
    _sub(doc, "Nhóm người dùng và ranh giới quyền")
    _bang(doc, ["Nhóm", "Việc chính", "Không được làm"],
          [[x.get("persona", ""), x.get("viec_chinh", ""), x.get("khong_duoc", "")] for x in d.get("persona", [])],
          rong=[3.4, 7.2, 5.9])

    # ══ 8 ══
    _muc(doc, 8, "Danh mục module", "Module catalogue")
    _bang(doc, ["Mã", "Module", "Mục đích", "Chủ quản"],
          [[x.get("ma", ""), x.get("ten", ""), x.get("muc_dich", ""), x.get("chu_quan", "")]
           for x in d.get("modules", [])], rong=[1.3, 4.4, 8.6, 2.2],
          mau_o=lambda ri, ci, v: GOLD_50 if (isinstance(v, str) and "LÕI" in v) else None)

    # ══ 9 ══
    yc = d.get("yeu_cau", [])
    _muc(doc, 9, "Yêu cầu nghiệp vụ chi tiết", "Detailed business requirements", ngat=True)
    _para(doc, (f"Tổng cộng {len(yc)} yêu cầu. " if yc else "") +
               "Mỗi yêu cầu có tiêu chí chấp nhận kiểm chứng được khi nghiệm thu. "
               "Danh mục dạng bảng lọc được nằm ở tệp Ma trận Yêu cầu kèm theo.", 9.8, italic=True, mau=INK_3)
    for m in d.get("modules", []):
        nhom = [y for y in yc if y.get("module") == m.get("ma")]
        if not nhom and not TRONG:
            continue
        _sub(doc, f"{m.get('ma')} · {m.get('ten')}")
        _bang(doc, ["Mã", "Yêu cầu", "Tiêu chí chấp nhận", "Ưu tiên", "Hiện trạng"],
              [[y.get("ma", ""), y.get("phat_bieu", ""), y.get("tieu_chi", ""),
                UU_TEN.get(y.get("uu_tien", ""), ""), y.get("trang_thai", "")] for y in nhom],
              rong=[2.1, 6.2, 5.4, 1.4, 1.4],
              mau_o=lambda ri, ci, v: TT_MAU.get(v) if ci == 4 else (RED_50 if (ci == 3 and v == "Bắt buộc") else None))

    # ══ 10 ══
    rg = d.get("ranh_gioi_ai", {})
    _muc(doc, 10, "Ranh giới AI & cổng người duyệt", "AI boundaries and human gates", ngat=True)
    _khung(doc, "Nguyên tắc bất biến",
           "AI hỗ trợ con người, không thay thế trách nhiệm con người. Danh sách dưới đây được cưỡng chế bằng "
           "mã nguồn ở tầng nền tảng, không phải bằng hướng dẫn sử dụng — nghĩa là không thể cấu hình để lách.",
           GREEN_50)
    _sub(doc, "AI ĐƯỢC PHÉP")
    _cham(doc, rg.get("duoc_phep", []), size=10)
    _sub(doc, "AI KHÔNG ĐƯỢC PHÉP")
    _cham(doc, rg.get("khong_duoc_phep", []), size=10)
    _sub(doc, "Bốn cổng người duyệt bắt buộc")
    _bang(doc, ["Cổng", "AI làm gì", "Người quyết định", "Bằng chứng lưu lại"],
          [[x.get("cong", ""), x.get("ai_lam", ""), x.get("nguoi_quyet", ""), x.get("bang_chung", "")]
           for x in rg.get("cong_hitl", [])], rong=[3.6, 4.6, 4.4, 3.9])

    # ══ 11 ══
    _muc(doc, 11, "Danh bạ AI agent", "AI agent registry")
    _para(doc, "Mỗi agent phải có danh tính khai báo trước khi được phép chạy. Agent không khai báo đủ các trường "
               "dưới đây thì không được kích hoạt ở môi trường thật.", 9.8, italic=True, mau=INK_3)
    _bang(doc, ["Agent", "Mục đích", "Dữ liệu sử dụng", "Phân loại tối đa", "Cổng người duyệt", "Chủ quản"],
          [[x.get("ten", ""), x.get("muc_dich", ""), x.get("du_lieu", ""), x.get("phan_loai_toi_da", ""),
            x.get("hitl", ""), x.get("chu_quan", "")] for x in d.get("agent", [])],
          rong=[2.9, 3.4, 3.2, 2.6, 2.9, 1.5])

    # ══ 12 ══
    _muc(doc, 12, "Dữ liệu & phân loại", "Data and classification", ngat=True)
    _para(doc, "iPMS không tạo kho dữ liệu song song. Dữ liệu do iPMS sinh ra được đăng ký và chia sẻ qua nền tảng "
               "dữ liệu tập đoàn; dữ liệu của hệ khác được tiêu thụ chứ không sao chép thành bản riêng.")
    _bang(doc, ["Nhóm dữ liệu", "Nguồn gốc", "Chủ dữ liệu", "Phân loại", "Ranh giới AI"],
          [[x.get("nhom", ""), x.get("nguon_goc", ""), x.get("chu_du_lieu", ""), x.get("phan_loai", ""),
            x.get("ranh_gioi_ai", "")] for x in d.get("du_lieu", [])], rong=[3.6, 3.2, 1.8, 2.6, 5.3])

    # ══ 13 ══
    _muc(doc, 13, "Yêu cầu phi chức năng", "Non-functional requirements")
    _bang(doc, ["Nhóm", "Yêu cầu", "Mục tiêu"],
          [[x.get("nhom", ""), x.get("yeu_cau", ""), x.get("muc_tieu", "")] for x in d.get("phi_chuc_nang", [])],
          rong=[3.0, 8.0, 5.5])

    # ══ 14 ══
    _muc(doc, 14, "Lộ trình & cổng ra", "Roadmap and exit gates", ngat=True)
    _para(doc, "Mỗi giai đoạn chỉ được coi là hoàn thành khi qua cổng ra — tiêu chí kiểm chứng được, "
               "không phải tự đánh giá.", 9.8, italic=True, mau=INK_3)
    _bang(doc, ["Giai đoạn", "Mục tiêu", "Kết quả", "Cổng ra"],
          [[x.get("giai_doan", ""), x.get("muc_tieu", ""), x.get("ket_qua", ""), x.get("cong_ra", "")]
           for x in d.get("lo_trinh", [])], rong=[3.4, 4.4, 4.0, 4.7])

    # ══ 15 ══
    _muc(doc, 15, "Rủi ro", "Risks")
    _bang(doc, ["Mã", "Loại", "Rủi ro", "Mức", "Giảm thiểu", "Trách nhiệm"],
          [[x.get("ma", ""), x.get("loai", ""), x.get("noi_dung", ""), x.get("muc", ""),
            x.get("giam_thieu", ""), x.get("trach_nhiem", "")] for x in d.get("rui_ro", [])],
          rong=[1.4, 1.9, 4.4, 1.4, 5.4, 2.0],
          mau_o=lambda ri, ci, v: RED_50 if (ci == 3 and v == "Cao") else None)

    # ══ 16 ══
    _muc(doc, 16, "Giả định", "Assumptions")
    _bang(doc, ["Mã", "Giả định", "Ảnh hưởng nếu sai"],
          [[x.get("ma", ""), x.get("noi_dung", ""), x.get("anh_huong", "")] for x in d.get("gia_dinh", [])],
          rong=[1.6, 7.6, 7.3])

    # ══ 17 ══
    _muc(doc, 17, "Câu hỏi mở", "Open questions")
    _bang(doc, ["#", "Câu hỏi cần quyết định", "Người quyết", "Chặn yêu cầu"],
          [[str(x.get("stt", "")), x.get("noi_dung", ""), x.get("nguoi_tra_loi", ""), x.get("anh_huong", "")]
           for x in d.get("cau_hoi_mo", [])], rong=[1.0, 9.5, 3.0, 3.0])

    # ══ 18 ══
    _muc(doc, 18, "Phê duyệt", "Approval")
    _para(doc, "Tài liệu có hiệu lực sau khi được phê duyệt. Thay đổi phạm vi sau ngày phê duyệt phải qua "
               "quy trình quản trị thay đổi và được ghi nhận bằng phiên bản mới.")
    t = doc.add_table(rows=2, cols=3); t.style = "Table Grid"
    for i, ten in enumerate(["CHỦ QUẢN NGHIỆP VỤ", "CHỦ QUẢN CÔNG NGHỆ", "CẤP PHÊ DUYỆT"]):
        c = t.rows[0].cells[i]; c.text = ""
        _txt(c.paragraphs[0], ten, 9.5, True, mau=GREEN_D)
        _to_cell(c, GREEN_50)
        c2 = t.rows[1].cells[i]; c2.text = ""
        for nhan in ["Chức danh: ..........................", "Ngày: ....................................",
                     "", "", "Ký và ghi rõ họ tên"]:
            _txt(c2.add_paragraph(), nhan, 9, mau=INK_3)

    ra.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(ra))


# ─────────────────────────── XLSX ───────────────────────────
VIEN = Border(*[Side(style="thin", color="DDE5E0")] * 4)


def _dau(ws, cot, rong, dong=1):
    f = PatternFill("solid", fgColor=GREEN)
    for i, (ten, w) in enumerate(zip(cot, rong), start=1):
        c = ws.cell(row=dong, column=i, value=ten)
        c.fill = f; c.font = Font(name=FONT, bold=True, size=10, color="FFFFFF")
        c.alignment = Alignment(vertical="center", wrap_text=True); c.border = VIEN
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.row_dimensions[dong].height = 30


def _o(ws, r, c, v, fill=None, bold=False, size=10, wrap=True):
    cell = ws.cell(row=r, column=c, value=v)
    cell.font = Font(name=FONT, size=size, bold=bold)
    cell.alignment = Alignment(vertical="top", wrap_text=wrap)
    cell.border = VIEN
    if fill:
        cell.fill = PatternFill("solid", fgColor=fill)
    return cell


def _sheet_bang(wb, ten, cot, rong, dong, mau_o=None, cao=None, khong_luoi=True, dong_trong=None):
    ws = wb.create_sheet(ten)
    if khong_luoi:
        ws.sheet_view.showGridLines = False
    if not dong and TRONG:                 # bảng chưa có dữ liệu → dòng trống để điền
        dong = [["" for _ in cot] for _ in range(dong_trong or SO_DONG_TRONG * 2)]
        mau_o = None
    _dau(ws, cot, rong)
    for i, r in enumerate(dong, start=2):
        for j, v in enumerate(r, start=1):
            _o(ws, i, j, v, fill=(mau_o(i - 2, j - 1, v) if mau_o else None))
        if cao:
            ws.row_dimensions[i].height = cao
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(cot))}{max(len(dong) + 1, 2)}"
    return ws


def dung_xlsx(d, ra: Path):
    wb = Workbook(); wb.remove(wb.active)
    tl, yc = d["tai_lieu"], d.get("yeu_cau", [])
    mod = {m.get("ma"): m.get("ten", "") for m in d.get("modules", [])}
    S1 = "01 · Danh mục Yêu cầu"

    # 00
    ws = wb.create_sheet("00 · Hướng dẫn đọc")
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 32; ws.column_dimensions["B"].width = 88
    _o(ws, 1, 1, "MA TRẬN YÊU CẦU — NỀN TẢNG iPMS (AI-NATIVE)" + (" · BẢN TRỐNG" if TRONG else ""),
       bold=True, size=15).font = Font(name=FONT, bold=True, size=15, color=GREEN_D)
    _o(ws, 2, 1, f"{tl.get('ma','')} · {tl.get('phien_ban','')} · {tl.get('ngay','')} · {tl.get('phan_loai','')}"
                 if not TRONG else "Điền thông tin tài liệu vào tệp .docx kèm theo trước khi dùng bảng này.")
    r = 4
    _o(ws, r, 1, "★ LÕI NỀN TẢNG", bold=True, fill=GOLD_50)
    _o(ws, r, 2, d.get("loi_nen_tang", {}).get("tuyen_bo", "")
       or "Xác định LÕI của nền tảng: năng lực mà mất nó thì sản phẩm mất lý do tồn tại. Ghi rõ ở Mục 4 của tệp .docx, "
          "rồi điền chuỗi truy vết vào sheet 03.", fill=GOLD_50)
    ws.row_dimensions[r].height = 46
    r += 2
    _o(ws, r, 1, "Sheet", bold=True, fill=GREEN_50); _o(ws, r, 2, "Nội dung", bold=True, fill=GREEN_50); r += 1
    for a, b in [
        (S1, "Toàn bộ yêu cầu nghiệp vụ, lọc được theo module, ưu tiên và hiện trạng"),
        ("02 · Tổng hợp theo module", "Đếm tự động theo sheet 01 — không nhập tay"),
        ("03 · Chuỗi truy vết LÕI", "★ Mắt xích từ chiến lược tới việc làm và ngược lại — trục xương sống của nền tảng"),
        ("04 · Ranh giới AI", "AI được phép và không được phép làm, bốn cổng người duyệt"),
        ("05 · Danh bạ AI agent", "Mỗi agent phải khai báo đủ trước khi được chạy"),
        ("06 · Dữ liệu & Phân loại", "Chủ dữ liệu, mức phân loại, ranh giới AI theo từng nhóm"),
        ("07 · Trung tâm quyền lực", "Quan hệ với V1, B0, B1, B2, B3, B5, B6, OpCo và dashboard cần có"),
        ("08 · Phi chức năng", "Quy mô, khả dụng, hạ tầng, định danh, ngôn ngữ, khả năng mở rộng"),
        ("09 · Lộ trình & Cổng ra", "Sáu giai đoạn kèm tiêu chí kiểm chứng được"),
        ("10 · Rủi ro & Giả định", "Rủi ro có người chịu trách nhiệm; giả định có ảnh hưởng nếu sai"),
        ("11 · 12 câu hỏi bắt buộc", "Phần soi chiếu bắt buộc trước khi phê duyệt xây dựng"),
        ("12 · Câu hỏi mở", "Điểm còn chờ quyết định và yêu cầu bị chặn tương ứng"),
    ]:
        _o(ws, r, 1, a, bold=True); _o(ws, r, 2, b); r += 1
    r += 1
    _o(ws, r, 1, "Chú giải hiện trạng", bold=True, fill=GREEN_50); _o(ws, r, 2, "", fill=GREEN_50); r += 1
    for k, v in [("Đã có", "Đã xây dựng và kiểm chứng"), ("Một phần", "Có nền, còn thiếu năng lực đã nêu"),
                 ("Chưa có", "Chưa xây dựng — cần đưa vào lộ trình")]:
        _o(ws, r, 1, k, bold=True, fill=TT_MAU[k]); _o(ws, r, 2, v); r += 1

    # 01
    _sheet_bang(wb, S1,
                ["Mã", "Module", "Yêu cầu", "Tiêu chí chấp nhận", "Ưu tiên", "Hiện trạng", "Lớp áp dụng", "Ghi chú"],
                [12, 26, 62, 54, 12, 13, 14, 30],
                [[y.get("ma", ""), f"{y.get('module','')} · {mod.get(y.get('module',''),'')}", y.get("phat_bieu", ""),
                  y.get("tieu_chi", ""), UU_TEN.get(y.get("uu_tien", ""), ""), y.get("trang_thai", ""),
                  y.get("lop", ""), y.get("ghi_chu", "")] for y in yc],
                mau_o=lambda ri, ci, v: (TT_MAU.get(v) if ci == 5 else
                                         (RED_50 if (ci == 4 and v == "Bắt buộc") else
                                          (GOLD_50 if (ci == 1 and "LÕI" in str(v)) else None))),
                cao=56, dong_trong=40)

    # 02
    ws = wb.create_sheet("02 · Tổng hợp theo module")
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 38
    for c in "BCDEFG":
        ws.column_dimensions[c].width = 14
    _o(ws, 1, 1, "TỔNG HỢP YÊU CẦU THEO MODULE", bold=True, size=13).font = \
        Font(name=FONT, bold=True, size=13, color=GREEN_D)
    _o(ws, 2, 1, "Các ô đếm tự động theo sheet 01.", size=9).font = Font(name=FONT, size=9, italic=True, color=INK_3)
    rng = f"'{S1}'"
    _dau(ws, ["Module", "Bắt buộc", "Nên có", "Đã có", "Một phần", "Chưa có", "Tổng"],
         [38, 14, 14, 14, 14, 14, 12], dong=4)
    for i, m in enumerate(d.get("modules", []), start=5):
        nhan = f"{m.get('ma')} · {m.get('ten')}"
        _o(ws, i, 1, nhan, bold=True, fill=GOLD_50 if "LÕI" in nhan else None)
        pre = f'{rng}!$B:$B,"{m.get("ma")} · *"'
        _o(ws, i, 2, f'=COUNTIFS({pre},{rng}!$E:$E,"Bắt buộc")', wrap=False)
        _o(ws, i, 3, f'=COUNTIFS({pre},{rng}!$E:$E,"Nên có")', wrap=False)
        _o(ws, i, 4, f'=COUNTIFS({pre},{rng}!$F:$F,"Đã có")', wrap=False, fill=GREEN_50)
        _o(ws, i, 5, f'=COUNTIFS({pre},{rng}!$F:$F,"Một phần")', wrap=False, fill=WARN_50)
        _o(ws, i, 6, f'=COUNTIFS({pre},{rng}!$F:$F,"Chưa có")', wrap=False, fill=RED_50)
        _o(ws, i, 7, f"=D{i}+E{i}+F{i}", wrap=False, bold=True)
    t = 5 + len(d.get("modules", []))
    _o(ws, t, 1, "TỔNG CỘNG", bold=True, fill=GREEN_50)
    for j in range(2, 8):
        col = get_column_letter(j)
        _o(ws, t, j, f"=SUM({col}5:{col}{t-1})", wrap=False, bold=True, fill=GREEN_50)
    _o(ws, t + 2, 1, "Yêu cầu Bắt buộc còn ở trạng thái Chưa có (cần đưa vào lộ trình ngay):", bold=True)
    _o(ws, t + 2, 2, f'=COUNTIFS({rng}!$E:$E,"Bắt buộc",{rng}!$F:$F,"Chưa có")', wrap=False, bold=True, fill=RED_50)

    # 03 — LÕI
    ln = d.get("loi_nen_tang", {})
    ws = wb.create_sheet("03 · Chuỗi truy vết LÕI")
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 26; ws.column_dimensions["B"].width = 46
    ws.column_dimensions["C"].width = 28; ws.column_dimensions["D"].width = 44
    _o(ws, 1, 1, "★ LÕI NỀN TẢNG — TỪ ĐIỂN TÁC VỤ GẮN KPI", bold=True, size=14).font = \
        Font(name=FONT, bold=True, size=14, color=GREEN_D)
    _o(ws, 2, 1, ln.get("tuyen_bo", ""), fill=GOLD_50)
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=4)
    ws.row_dimensions[2].height = 40
    _dau(ws, ["Mắt xích", "Câu hỏi nó trả lời", "Chủ thể", "Ví dụ"], [26, 46, 28, 44], dong=4)
    for i, x in enumerate(ln.get("chuoi_truy_vet", []), start=5):
        _o(ws, i, 1, x.get("mat_xich", ""), bold=True, fill=GREEN_50)
        _o(ws, i, 2, x.get("cau_hoi", ""))
        _o(ws, i, 3, x.get("chu_the", ""))
        _o(ws, i, 4, x.get("vi_du", ""))
        ws.row_dimensions[i].height = 32
    r = 5 + len(ln.get("chuoi_truy_vet", [])) + 1
    _o(ws, r, 1, "Ràng buộc bất biến", bold=True, fill=RED_50)
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=4); r += 1
    for x in ln.get("rang_buoc_bat_bien", []):
        _o(ws, r, 1, "• " + x)
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=4)
        ws.row_dimensions[r].height = 26; r += 1
    r += 1
    _dau(ws, ["Lõi này là nền để triển khai AI", "Ý nghĩa", "", ""], [26, 46, 28, 44], dong=r); r += 1
    for x in ln.get("vai_tro_ai", []):
        _o(ws, r, 1, x.get("nang_luc", ""), bold=True)
        _o(ws, r, 2, x.get("y_nghia", ""))
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=4)
        ws.row_dimensions[r].height = 30; r += 1

    # 04
    rg = d.get("ranh_gioi_ai", {})
    ws = wb.create_sheet("04 · Ranh giới AI")
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 46; ws.column_dimensions["B"].width = 46
    _dau(ws, ["AI ĐƯỢC PHÉP", "AI KHÔNG ĐƯỢC PHÉP"], [46, 46])
    dp, kdp = rg.get("duoc_phep", []), rg.get("khong_duoc_phep", [])
    for i in range(max(len(dp), len(kdp))):
        _o(ws, i + 2, 1, dp[i] if i < len(dp) else "", fill=GREEN_50 if i < len(dp) else None)
        _o(ws, i + 2, 2, kdp[i] if i < len(kdp) else "", fill=RED_50 if i < len(kdp) else None)
    r = max(len(dp), len(kdp)) + 3
    _dau(ws, ["Cổng người duyệt", "AI làm gì", "Người quyết định", "Bằng chứng lưu lại"], [30, 40, 34, 34], dong=r)
    for i, x in enumerate(rg.get("cong_hitl", []), start=r + 1):
        _o(ws, i, 1, x.get("cong", ""), bold=True, fill=GREEN_50)
        _o(ws, i, 2, x.get("ai_lam", "")); _o(ws, i, 3, x.get("nguoi_quyet", "")); _o(ws, i, 4, x.get("bang_chung", ""))

    # 05–12
    _sheet_bang(wb, "05 · Danh bạ AI agent",
                ["Agent", "Mục đích", "Dữ liệu sử dụng", "Phân loại tối đa", "Cổng người duyệt", "Chủ quản"],
                [28, 34, 34, 30, 34, 14],
                [[x.get("ten", ""), x.get("muc_dich", ""), x.get("du_lieu", ""), x.get("phan_loai_toi_da", ""),
                  x.get("hitl", ""), x.get("chu_quan", "")] for x in d.get("agent", [])], cao=34)
    _sheet_bang(wb, "06 · Dữ liệu & Phân loại",
                ["Nhóm dữ liệu", "Nguồn gốc", "Chủ dữ liệu", "Phân loại", "Ranh giới AI"],
                [38, 30, 16, 26, 48],
                [[x.get("nhom", ""), x.get("nguon_goc", ""), x.get("chu_du_lieu", ""), x.get("phan_loai", ""),
                  x.get("ranh_gioi_ai", "")] for x in d.get("du_lieu", [])],
                mau_o=lambda ri, ci, v: (RED_50 if (ci == 3 and "Restricted" in str(v)) else
                                         (WARN_50 if (ci == 3 and "Confidential" in str(v)) else None)), cao=34)
    _sheet_bang(wb, "07 · Trung tâm quyền lực",
                ["Mã", "Khối", "Vai trò", "Quan hệ với iPMS", "Dashboard cần có"],
                [8, 32, 32, 44, 48],
                [[x.get("ma", ""), x.get("ten", ""), x.get("vai_tro", ""), x.get("quan_he_ipms", ""),
                  x.get("dashboard", "")] for x in d.get("trung_tam_quyen_luc", [])], cao=34)
    _sheet_bang(wb, "08 · Phi chức năng", ["Nhóm", "Yêu cầu", "Mục tiêu"], [24, 62, 44],
                [[x.get("nhom", ""), x.get("yeu_cau", ""), x.get("muc_tieu", "")] for x in d.get("phi_chuc_nang", [])],
                cao=32)
    _sheet_bang(wb, "09 · Lộ trình & Cổng ra", ["Giai đoạn", "Mục tiêu", "Kết quả", "Cổng ra (kiểm chứng được)"],
                [28, 40, 36, 48],
                [[x.get("giai_doan", ""), x.get("muc_tieu", ""), x.get("ket_qua", ""), x.get("cong_ra", "")]
                 for x in d.get("lo_trinh", [])], cao=38)

    ws = _sheet_bang(wb, "10 · Rủi ro & Giả định", ["Mã", "Loại", "Nội dung", "Mức", "Giảm thiểu", "Trách nhiệm"],
                     [10, 16, 46, 12, 52, 18],
                     [[x.get("ma", ""), x.get("loai", ""), x.get("noi_dung", ""), x.get("muc", ""),
                       x.get("giam_thieu", ""), x.get("trach_nhiem", "")] for x in d.get("rui_ro", [])],
                     mau_o=lambda ri, ci, v: RED_50 if (ci == 3 and v == "Cao") else None, cao=40)
    r = len(d.get("rui_ro", [])) + 3
    _dau(ws, ["Mã", "Giả định", "Ảnh hưởng nếu sai", "", "", ""], [10, 16, 46, 12, 52, 18], dong=r)
    for i, x in enumerate(d.get("gia_dinh", []), start=r + 1):
        _o(ws, i, 1, x.get("ma", ""), bold=True)
        _o(ws, i, 2, x.get("noi_dung", "")); ws.merge_cells(start_row=i, start_column=2, end_row=i, end_column=3)
        _o(ws, i, 4, x.get("anh_huong", "")); ws.merge_cells(start_row=i, start_column=4, end_row=i, end_column=6)
        ws.row_dimensions[i].height = 34

    _sheet_bang(wb, "11 · 12 câu hỏi bắt buộc", ["#", "Câu hỏi", "Trả lời của iPMS"], [6, 44, 88],
                [[x.get("stt", ""), x.get("cau_hoi", ""), x.get("tra_loi", "")] for x in d.get("muoi_hai_cau_hoi", [])],
                cao=52)
    _sheet_bang(wb, "12 · Câu hỏi mở", ["#", "Câu hỏi cần quyết định", "Người quyết", "Chặn yêu cầu", "Ngày cần có"],
                [6, 62, 20, 24, 20],
                [[x.get("stt", ""), x.get("noi_dung", ""), x.get("nguoi_tra_loi", ""), x.get("anh_huong", ""), ""]
                 for x in d.get("cau_hoi_mo", [])], cao=34)

    ra.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(ra))


def main():
    try:
        sys.stdout.reconfigure(encoding="utf-8"); sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass
    ap = argparse.ArgumentParser(description="Sinh BRD Nền tảng iPMS (AI-Native)")
    ap.add_argument("nguon"); ap.add_argument("--out", default=None)
    ap.add_argument("--trong", action="store_true",
                    help="sinh BẢN TRỐNG để điền: giữ nguyên khung 18 mục, xoá mọi câu trả lời, "
                         "thêm dòng hướng dẫn và ô/dòng trống")
    ap.add_argument("--dong-trong", type=int, default=6, help="số dòng trống mỗi bảng (mặc định 6)")
    a = ap.parse_args()

    global TRONG, SO_DONG_TRONG
    TRONG = a.trong
    SO_DONG_TRONG = a.dong_trong

    src = Path(a.nguon).resolve()
    d = bo_khoa_soan_thao(json.loads(src.read_text(encoding="utf-8")))
    if TRONG:
        d = lam_trong(d)
    out = Path(a.out).resolve() if a.out else src.parent
    pb = d.get("tai_lieu", {}).get("phien_ban", "") or "v1"
    hau_to = "_BAN_TRONG" if TRONG else f"_{pb}"
    f_doc = out / f"01_BRD_iPMS_AI_Native{hau_to}.docx"
    f_xls = out / f"02_Ma_tran_Yeu_cau_iPMS{hau_to}.xlsx"

    dung_docx(d, f_doc)
    dung_xlsx(d, f_xls)

    yc = d.get("yeu_cau", [])
    print(("📄 BẢN TRỐNG — để điền" if TRONG else "📄 Bản đã điền"))
    print(f"✅ {f_doc.name}   ({f_doc.stat().st_size:,} bytes)")
    print(f"✅ {f_xls.name}   ({f_xls.stat().st_size:,} bytes)")
    print(f"   Thư mục: {out}")
    if TRONG:
        print(f"   Khung giữ nguyên: 18 mục · {len(d.get('modules', []))} module · "
              f"{len(d.get('muoi_hai_cau_hoi', []))} câu hỏi bắt buộc · "
              f"{len(d.get('loi_nen_tang', {}).get('chuoi_truy_vet', []))} mắt xích chuỗi truy vết")
    else:
        print(f"   {len(yc)} yêu cầu / {len(d.get('modules', []))} module · "
              f"Bắt buộc: {sum(1 for y in yc if y.get('uu_tien') == 'M')} · "
              f"Chưa có: {sum(1 for y in yc if y.get('trang_thai') == 'Chưa có')}")


if __name__ == "__main__":
    main()
