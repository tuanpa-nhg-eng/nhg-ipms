#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Sinh BỘ HỒ SƠ GỬI KHÁCH từ một nguồn duy nhất (ho-so.json).

    py xuat-ho-so.py <đường-dẫn-ho-so.json> [--out <thư-mục>]

Đầu ra:
    01_BRD_{MÃ-KH}_{phiên bản}.docx      — tài liệu chính, khách rà bằng Track Changes
    02_Phu_luc_{MÃ-KH}_{phiên bản}.xlsx  — 7 sheet, khách ghi ý kiến trực tiếp

HÀNG RÀO NỘI BỘ nằm trong chính bộ sinh file:
  · Mọi khoá bắt đầu bằng "_" bị loại bỏ đệ quy TRƯỚC khi dựng tài liệu.
  · Sau khi loại, quét mẫu dấu vết nội bộ; dính một mẫu là DỪNG, không ghi file.
Không có đường nào để ghi chú nội bộ chảy ra bản gửi khách.
"""
import argparse, json, re, sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation

ROOT = Path(__file__).resolve().parents[4]
LOGO = ROOT / "design-system" / "public" / "logo.png"

GREEN, GREEN_D, RED, INK, INK_3 = "037236", "04361C", "ED2024", "101214", "7C8285"
GREEN_50, MIST, WARN_50, INFO_50, RED_50 = "E8F4ED", "EEF3EF", "FBF0DC", "E5F0F9", "FDECEC"
FONT = "Be Vietnam Pro"

# Tầng 1 — CHẶN CỨNG: những chuỗi không thể xuất hiện hợp lệ trong tài liệu gửi khách.
CAM = re.compile(
    r"\bF[0-9]{2,3}\b|OWNER_DIGEST|STATUS\.md|Reviewer Agent|reviewer đối kháng"
    r"|\bcommit\b|[0-9]+/[0-9]+ (PASS|test)|OneOffice|Bravo|\bH\.01\b|\bOpCo\b"
    r"|Nguyễn Hoàng|ngân sách nội bộ|\btrđ\b",
    re.IGNORECASE,
)

# Tầng 2 — CẢNH BÁO: từ nhập nhằng. Hợp lệ khi mô tả điểm đau của khách
# ("mất 3–5 ngày công mỗi kỳ"), KHÔNG hợp lệ khi là ước lượng nguồn lực của mình.
# Chỉ nhắc để người thật liếc mắt, không chặn — chặn cứng ở đây sẽ dạy người dùng
# thói quen bỏ qua cảnh báo, và đó là cách một hàng rào chết.
CANH_BAO = re.compile(r"ngày công|đơn giá|báo giá|chi phí triển khai|nhân sự dự án", re.IGNORECASE)

FIT_MAU = {
    "Sẵn sàng": GREEN_50,
    "Sẵn sàng — cần kích hoạt": INFO_50,
    "Đáp ứng bằng cấu hình": MIST,
    "Cần phát triển": WARN_50,
    "Ngoài phạm vi đợt này": RED_50,
}
PHAN_HE = {
    "A": "Động cơ & kỳ vọng", "B": "Phân rã mục tiêu", "C": "Cơ cấu tổ chức",
    "D": "Vai trò & phân quyền", "E": "Khung KPI", "F": "Chu trình đánh giá",
    "G": "Kiến trúc tác vụ", "H": "Dữ liệu & tích hợp", "I": "AI & tự động hoá",
    "J": "Báo cáo", "K": "Cấu hình & thương hiệu", "L": "Phi chức năng & tuân thủ",
    "M": "Phạm vi & lộ trình",
}
MOSCOW = {"M": "Bắt buộc", "S": "Nên có", "C": "Có thì tốt", "W": "Đợt sau"}


# ─────────────────────────── hàng rào ───────────────────────────
def bo_khoa_noi_bo(x):
    """Loại đệ quy mọi khoá bắt đầu bằng '_'."""
    if isinstance(x, dict):
        return {k: bo_khoa_noi_bo(v) for k, v in x.items() if not k.startswith("_")}
    if isinstance(x, list):
        return [bo_khoa_noi_bo(v) for v in x]
    return x


def quet_hang_rao(x, duong="", loi=None, canh=None):
    loi = [] if loi is None else loi
    canh = [] if canh is None else canh
    if isinstance(x, dict):
        for k, v in x.items():
            quet_hang_rao(v, f"{duong}.{k}", loi, canh)
    elif isinstance(x, list):
        for i, v in enumerate(x):
            quet_hang_rao(v, f"{duong}[{i}]", loi, canh)
    elif isinstance(x, str):
        m = CAM.search(x)
        if m:
            loi.append((duong, m.group(0), x[:90]))
        w = CANH_BAO.search(x)
        if w:
            canh.append((duong, w.group(0), x[:90]))
    return loi, canh


# ─────────────────────────── DOCX ───────────────────────────
def _field(par, ma):
    r = par.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    t = OxmlElement("w:instrText"); t.set(qn("xml:space"), "preserve"); t.text = ma
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    r._r.append(b); r._r.append(t); r._r.append(e)
    return r


def _to_cell(cell, mau):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), mau)
    tcPr.append(shd)


def _txt(par, text, size=10.5, bold=False, italic=False, mau=INK):
    r = par.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = RGBColor.from_string(mau)
    return r


def _muc(doc, so, vi, en):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    _txt(p, f"{so}. ", 14, True, mau=RED)
    _txt(p, vi, 14, True, mau=GREEN_D)
    if en:
        p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.keep_with_next = True
        _txt(p2, en, 9, italic=True, mau=INK_3)
    return p


def _sub(doc, vi, en=""):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    _txt(p, vi, 11.5, True, mau=GREEN)
    if en:
        _txt(p, "  " + en, 9, italic=True, mau=INK_3)


def _para(doc, text, size=10.5, italic=False, mau=INK, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _txt(p, text, size, italic=italic, mau=mau)
    return p


def _bang(doc, dau_cot, dong, rong=None):
    t = doc.add_table(rows=1, cols=len(dau_cot)); t.style = "Table Grid"
    t.autofit = True
    for i, h in enumerate(dau_cot):
        c = t.rows[0].cells[i]; c.text = ""
        _txt(c.paragraphs[0], h, 9.5, True, mau="FFFFFF")
        _to_cell(c, GREEN)
    for r in dong:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = ""
            _txt(cells[i].paragraphs[0], "" if v is None else str(v), 9.5)
    if rong:
        for i, w in enumerate(rong):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _the_yeu_cau(doc, y):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    _txt(p, y.get("ma", ""), 10.5, True, mau=GREEN)
    _txt(p, "  " + y.get("tieu_de", ""), 11, True)
    uu = y.get("moscow", "")
    if uu:
        _txt(p, f"   [{MOSCOW.get(uu, uu)}]", 8.5, True, mau=RED if uu == "M" else INK_3)
    hang = [
        ("Yêu cầu", y.get("phat_bieu", "")),
        ("Vì sao", y.get("vi_sao", "")),
        ("Tiêu chí chấp nhận", y.get("tieu_chi_chap_nhan", "")),
        ("Mức đáp ứng", y.get("muc_fit", "")),
    ]
    if y.get("gap"):
        hang.append(("Khoảng trống", y["gap"]))
    ng = " · ".join(x for x in [y.get("nguon", ""), y.get("bang_chung", "")] if x)
    if ng:
        hang.append(("Nguồn", ng))
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
    for k, v in hang:
        c = t.add_row().cells
        c[0].text = ""; _txt(c[0].paragraphs[0], k, 9, True, mau=GREEN)
        _to_cell(c[0], GREEN_50)
        c[0].width = Cm(3.6)
        c[1].text = ""; _txt(c[1].paragraphs[0], str(v), 9.5)
        c[1].width = Cm(12.9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def dung_docx(d, ra: Path):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT; st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2)

    # chân trang: phân loại + số trang
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _txt(fp, "MẬT — chỉ lưu hành giữa hai bên  ·  Trang ", 8, mau=INK_3)
    _field(fp, "PAGE"); _txt(fp, " / ", 8, mau=INK_3); _field(fp, "NUMPAGES")
    for r in fp.runs:
        r.font.size = Pt(8); r.font.name = FONT; r.font.color.rgb = RGBColor.from_string(INK_3)

    kh, tl = d["khach"], d["tai_lieu"]

    # ── trang bìa ──
    if LOGO.exists():
        pl = doc.add_paragraph(); pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pl.add_run().add_picture(str(LOGO), height=Cm(1.2))
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(60)
    _txt(p, "TÀI LIỆU YÊU CẦU NGHIỆP VỤ", 26, True, mau=GREEN_D)
    p = doc.add_paragraph(); _txt(p, "Business Requirements Document", 12, italic=True, mau=INK_3)
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(18)
    _txt(p, kh.get("ten", ""), 16, True)
    p = doc.add_paragraph(); _txt(p, "Triển khai Nền tảng Quản trị Hiệu suất iPMS", 12, mau=INK)
    p = doc.add_paragraph(); _txt(p, f"iPMS Platform implementation — {kh.get('ten_en','')}", 9.5, italic=True, mau=INK_3)

    doc.add_paragraph().paragraph_format.space_before = Pt(24)
    _bang(doc, ["", ""], [
        ["Mã hồ sơ", d.get("ma_kh", "")],
        ["Phiên bản", f"{tl.get('phien_ban','')} · {tl.get('ngay','')}"],
        ["Phân loại", "MẬT — chỉ lưu hành giữa hai bên"],
        ["Số yêu cầu", str(len(d.get("yeu_cau", [])))],
        ["Căn cứ", tl.get("can_cu", "")],
    ], rong=[4.5, 12.0])
    _para(doc, "Tài liệu ghi nhận yêu cầu nghiệp vụ của Quý tổ chức và mức đáp ứng của Nền tảng iPMS, "
               "làm căn cứ thống nhất phạm vi trước khi lập Đề xuất giải pháp.", 9.5, italic=True, mau=INK_3)
    _para(doc, "Kèm theo: Phụ lục Ma trận Yêu cầu (tệp bảng tính) — Quý vị ghi ý kiến trực tiếp vào cột dành riêng.",
          9.5, italic=True, mau=INK_3)
    doc.add_page_break()

    # ── 1 ──
    _muc(doc, 1, "Thông tin tài liệu", "Document control")
    _bang(doc, ["Nội dung", "Chi tiết"], [
        ["Khách hàng", kh.get("ten", "")],
        ["Lĩnh vực · Quy mô", f"{kh.get('nganh','')} · {kh.get('quy_mo','')} · {kh.get('so_don_vi','')}"],
        ["Dự án", f"Triển khai Nền tảng iPMS — {tl.get('pham_vi_ngan','')}"],
        ["Phiên bản", f"{tl.get('phien_ban','')} · {tl.get('ngay','')}"],
        ["Người lập", tl.get("nguoi_lap", "")],
        ["Người soát xét phía khách hàng", tl.get("nguoi_soat_xet_khach", "")],
    ], rong=[5.5, 11.0])

    # ── 2 ──
    _muc(doc, 2, "Bối cảnh & động cơ", "Context and drivers")
    bc = d.get("boi_canh", {})
    _para(doc, bc.get("tom_tat", ""))
    _sub(doc, "Cách làm hiện tại và chỗ nó gãy", "Current approach and where it breaks")
    _bang(doc, ["Điểm đau", "Hệ quả", "Bằng chứng"],
          [[x.get("noi_dung", ""), x.get("he_qua", ""), x.get("bang_chung", "")] for x in bc.get("diem_dau", [])],
          rong=[8.0, 6.0, 2.5])
    _para(doc, f"Mức trưởng thành hiện tại: {bc.get('muc_truong_thanh','')} — {bc.get('dien_giai','')}", 10)

    # ── 3 ──
    _muc(doc, 3, "Mục tiêu kinh doanh & tiêu chí thành công", "Business objectives and success criteria")
    _bang(doc, ["Mục tiêu kinh doanh", "Tiêu chí thành công đo được", "Ai đo", "Khi nào"],
          [[x.get("muc_tieu", ""), x.get("tieu_chi", ""), x.get("ai_do", ""), x.get("khi_nao", "")]
           for x in d.get("muc_tieu", [])], rong=[4.8, 6.2, 3.0, 2.5])

    # ── 4 ──
    _muc(doc, 4, "Phạm vi", "Scope")
    pv = d.get("pham_vi", {})
    _sub(doc, "Trong phạm vi đợt này", "In scope")
    _bang(doc, ["Hạng mục", "Đơn vị áp dụng", "Ghi chú"],
          [[x.get("hang_muc", ""), x.get("don_vi", ""), x.get("ghi_chu", "")] for x in pv.get("trong", [])],
          rong=[8.5, 4.0, 4.0])
    _sub(doc, "Ngoài phạm vi đợt này — hai bên đã thống nhất", "Out of scope — mutually agreed")
    _bang(doc, ["Hạng mục", "Lý do hoãn", "Xem xét lại"],
          [[x.get("hang_muc", ""), x.get("ly_do", ""), x.get("xem_xet", "")] for x in pv.get("ngoai", [])],
          rong=[6.0, 6.0, 4.5])

    # ── 5 ──
    _muc(doc, 5, "Các bên liên quan & nhóm người dùng", "Stakeholders and user groups")
    _bang(doc, ["Nhóm", "Quy mô", "Việc chính", "Được thấy", "Không được thấy"],
          [[x.get("nhom", ""), x.get("quy_mo", ""), x.get("viec_chinh", ""),
            x.get("duoc_thay", ""), x.get("khong_duoc_thay", "")] for x in d.get("nhom_nguoi_dung", [])],
          rong=[2.8, 1.8, 4.6, 3.6, 3.7])

    # ── 6 ──
    _muc(doc, 6, "Hiện trạng", "Current state (as-is)")
    _bang(doc, ["Phân hệ", "Hiện trạng"],
          [[f"{k} · {PHAN_HE.get(k,'')}", v] for k, v in d.get("as_is", {}).items() if v],
          rong=[4.5, 12.0])

    # ── 7 ──
    yc = d.get("yeu_cau", [])
    _muc(doc, 7, "Yêu cầu nghiệp vụ", "Business requirements")
    _para(doc, f"Tổng cộng {len(yc)} yêu cầu. Danh mục đầy đủ dạng bảng lọc được nằm ở Phụ lục "
               f"— Quý vị ghi ý kiến trực tiếp vào cột “Ý kiến của Quý vị”.", 9.5, italic=True, mau=INK_3)
    for ph in sorted({y.get("phan_he", "") for y in yc}):
        nhom = [y for y in yc if y.get("phan_he") == ph]
        if not nhom:
            continue
        _sub(doc, f"{ph} · {PHAN_HE.get(ph, '')}")
        for y in nhom:
            _the_yeu_cau(doc, y)

    # ── 8 ──
    _muc(doc, 8, "Hệ nguồn dữ liệu & tích hợp", "Data sources and integration")
    _bang(doc, ["Hệ nguồn", "Chủ sở hữu", "Dữ liệu cần lấy", "Có kết nối?", "Tần suất", "Dự phòng"],
          [[x.get("he", ""), x.get("chu_so_huu", ""), x.get("du_lieu", ""), x.get("co_api", ""),
            x.get("tan_suat", ""), x.get("du_phong", "")] for x in d.get("he_nguon", [])],
          rong=[3.0, 2.8, 3.8, 2.0, 2.2, 2.7])

    # ── 9 · TỪ ĐIỂN TÁC VỤ ──
    ktv = d.get("kien_truc_tac_vu", {})
    _muc(doc, 9, "Kiến trúc tác vụ & Từ điển Tác vụ", "Task architecture and task dictionary")
    _para(doc, ktv.get("ghi_chu", ""), 10)
    _para(doc, "Từ điển Tác vụ là nơi công việc thực tế của từng vị trí được mô tả một cách chuẩn hoá và "
               "gắn với chỉ số đo lường. Đây là mắt xích nối giữa mục tiêu và việc làm hằng ngày: mục tiêu "
               "phân rã thành chỉ số, chỉ số gắn vào tác vụ, tác vụ thuộc về một vị trí cụ thể. Nhờ vậy, "
               "kết quả đánh giá luôn giải thích được bằng công việc có thật.", 10)
    _para(doc, "Danh mục tác vụ dưới đây là bản khởi tạo. Sau khi hệ thống vận hành, chính trưởng phòng và "
               "nhân viên là người tiếp tục hoàn thiện theo vòng lặp: góp ý từ người làm thật → mở lại để sửa "
               "→ trưởng phòng duyệt → phiên bản mới, giữ nguyên lịch sử.", 10)
    _bang(doc, ["Phòng", "Nhóm tác vụ", "Tên tác vụ", "Người chịu trách nhiệm", "Đầu ra", "Mã chỉ số"],
          [[x.get("phong", ""), x.get("nhom_tac_vu", ""), x.get("ten_tac_vu", ""),
            x.get("nguoi_chiu_trach_nhiem", ""), x.get("dau_ra", ""), x.get("ma_kpi", "")]
           for x in ktv.get("dong", [])], rong=[2.4, 2.8, 4.2, 2.8, 2.6, 1.7])
    _para(doc, "Bảng đầy đủ để Quý vị bổ sung nằm ở Phụ lục, sheet “04 · Kiến trúc Tác vụ”.",
          9.5, italic=True, mau=INK_3)

    # ── 10 ──
    ai = d.get("ai", {})
    _muc(doc, 10, "Yêu cầu về trí tuệ nhân tạo & quản trị AI", "AI and AI governance")
    _para(doc, "Nguyên tắc bất biến của nền tảng: trí tuệ nhân tạo chỉ đề xuất, người có thẩm quyền là người "
               "quyết định; mọi đề xuất đều nêu được căn cứ và được lưu vết.", 10)
    _bang(doc, ["Nội dung", "Yêu cầu của Quý tổ chức"], [
        ["Kỳ vọng ứng dụng", ai.get("ky_vong", "")],
        ["Mức tự chủ chấp nhận được", ai.get("tu_chu", "")],
        ["Ranh giới xử lý dữ liệu", ai.get("egress", "")],
        ["Yêu cầu giải thích được", ai.get("giai_thich", "")],
        ["Chính sách chi phí", ai.get("chi_phi", "")],
    ], rong=[5.0, 11.5])

    # ── 11 ──
    _muc(doc, 11, "Yêu cầu báo cáo & phân tích", "Reporting and analytics")
    _bang(doc, ["Nhóm người dùng", "Báo cáo cần có", "Tần suất", "Định dạng"],
          [[x.get("nhom", ""), x.get("bao_cao", ""), x.get("tan_suat", ""), x.get("dinh_dang", "")]
           for x in d.get("bao_cao", [])], rong=[3.5, 6.5, 3.0, 3.5])

    # ── 12 ──
    _muc(doc, 12, "Yêu cầu phi chức năng, bảo mật & tuân thủ", "Non-functional, security and compliance")
    _bang(doc, ["Nhóm", "Yêu cầu", "Ghi chú"],
          [[x.get("nhom", ""), x.get("yeu_cau", ""), x.get("ghi_chu", "")] for x in d.get("nfr", [])],
          rong=[4.0, 8.5, 4.0])

    # ── 13 ──
    _muc(doc, 13, "Đối chiếu năng lực & khoảng trống", "Capability fit and gaps")
    dem = {k: 0 for k in FIT_MAU}
    dem_m = {k: 0 for k in FIT_MAU}
    for y in yc:
        f = y.get("muc_fit", "")
        if f in dem:
            dem[f] += 1
            if y.get("moscow") == "M":
                dem_m[f] += 1
    _bang(doc, ["Mức đáp ứng", "Số yêu cầu", "Trong đó bắt buộc"],
          [[k, str(dem[k]), str(dem_m[k])] for k in FIT_MAU], rong=[8.0, 4.0, 4.5])
    _para(doc, f"Đối chiếu với năng lực nền tảng tại ngày {tl.get('ngay_doi_chieu_nang_luc','')}.",
          9.5, italic=True, mau=INK_3)
    gaps = [y for y in yc if y.get("gap")]
    if gaps:
        _sub(doc, "Khoảng trống cần xử lý", "Gaps to resolve")
        _bang(doc, ["Yêu cầu", "Khoảng trống", "Mức đáp ứng"],
              [[y.get("ma", ""), y.get("gap", ""), y.get("muc_fit", "")] for y in gaps],
              rong=[2.5, 10.0, 4.0])
    _para(doc, "Nguồn lực và chi phí cho các hạng mục “Cần phát triển” được trình bày trong Đề xuất giải pháp "
               "riêng, sau khi hai bên thống nhất tài liệu này.", 9.5, italic=True, mau=INK_3)

    # ── 14 ──
    _muc(doc, 14, "Giả định · Ràng buộc · Phụ thuộc · Rủi ro", "Assumptions, constraints, dependencies and risks")
    _sub(doc, "Giả định", "Assumptions")
    _bang(doc, ["#", "Giả định", "Ảnh hưởng nếu sai"],
          [[x.get("ma", ""), x.get("noi_dung", ""), x.get("anh_huong", "")] for x in d.get("gia_dinh", [])],
          rong=[1.8, 8.5, 6.2])
    _sub(doc, "Ràng buộc", "Constraints")
    _bang(doc, ["#", "Ràng buộc", "Nguồn"],
          [[x.get("ma", ""), x.get("noi_dung", ""), x.get("nguon", "")] for x in d.get("rang_buoc", [])],
          rong=[1.8, 10.2, 4.5])
    _sub(doc, "Phụ thuộc phía Quý tổ chức", "Customer-side dependencies")
    _bang(doc, ["#", "Quý tổ chức cần cung cấp", "Phục vụ yêu cầu", "Hạn đề nghị"],
          [[x.get("ma", ""), x.get("noi_dung", ""), x.get("cho_yeu_cau", ""), x.get("han", "")]
           for x in d.get("phu_thuoc", [])], rong=[1.8, 8.0, 3.2, 3.5])
    _sub(doc, "Rủi ro", "Risks")
    _bang(doc, ["#", "Rủi ro", "Mức", "Giảm thiểu", "Trách nhiệm"],
          [[x.get("ma", ""), x.get("noi_dung", ""), x.get("muc", ""), x.get("giam_thieu", ""), x.get("trach_nhiem", "")]
           for x in d.get("rui_ro", [])], rong=[1.6, 5.0, 1.8, 5.6, 2.5])

    # ── 15 ──
    _muc(doc, 15, "Câu hỏi mở & bước tiếp theo", "Open questions and next steps")
    _bang(doc, ["#", "Câu hỏi cần Quý vị làm rõ", "Người trả lời", "Hạn", "Chặn yêu cầu"],
          [[str(x.get("stt", "")), x.get("noi_dung", ""), x.get("nguoi_tra_loi", ""), x.get("han", ""),
            x.get("chan_yeu_cau", "")] for x in d.get("cau_hoi_mo", [])], rong=[1.2, 7.4, 3.2, 2.4, 2.3])
    _para(doc, d.get("buoc_tiep_theo", ""), before=6)

    # ── ký ──
    doc.add_paragraph().paragraph_format.space_before = Pt(10)
    _para(doc, "Hai bên xác nhận nội dung tài liệu này phản ánh đúng yêu cầu nghiệp vụ tại thời điểm ký. "
               "Mọi thay đổi sau ngày ký được xử lý qua quy trình quản trị thay đổi phạm vi.", 10)
    t = doc.add_table(rows=2, cols=2); t.style = "Table Grid"
    for i, ten in enumerate(["ĐẠI DIỆN KHÁCH HÀNG", "ĐẠI DIỆN NHÀ CUNG CẤP"]):
        c = t.rows[0].cells[i]; c.text = ""
        _txt(c.paragraphs[0], ten, 10, True, mau=GREEN_D)
        _to_cell(c, GREEN_50)
        c2 = t.rows[1].cells[i]; c2.text = ""
        for nhan in ["Chức danh: ...........................................",
                     "Ngày: ....................................................",
                     "", "", "Ký và ghi rõ họ tên"]:
            pp = c2.add_paragraph(); _txt(pp, nhan, 9.5, mau=INK_3)
    ra.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(ra))


# ─────────────────────────── XLSX ───────────────────────────
VIEN = Border(*[Side(style="thin", color="DDE5E0")] * 4)


def _dau(ws, cot, rong, dong=1):
    f = PatternFill("solid", fgColor=GREEN)
    for i, (ten, w) in enumerate(zip(cot, rong), start=1):
        c = ws.cell(row=dong, column=i, value=ten)
        c.fill = f; c.font = Font(name=FONT, bold=True, size=10, color="FFFFFF")
        c.alignment = Alignment(vertical="center", wrap_text=True)
        c.border = VIEN
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.row_dimensions[dong].height = 30


def _o(ws, r, c, v, wrap=True, fill=None, bold=False, size=10):
    cell = ws.cell(row=r, column=c, value=v)
    cell.font = Font(name=FONT, size=size, bold=bold)
    cell.alignment = Alignment(vertical="top", wrap_text=wrap)
    cell.border = VIEN
    if fill:
        cell.fill = PatternFill("solid", fgColor=fill)
    return cell


def dung_xlsx(d, ra: Path):
    wb = Workbook(); wb.remove(wb.active)
    yc = d.get("yeu_cau", [])
    kh, tl = d["khach"], d["tai_lieu"]

    # ── 00 Hướng dẫn ──
    ws = wb.create_sheet("00 · Hướng dẫn đọc")
    ws.sheet_view.showGridLines = False
    _o(ws, 1, 1, "PHỤ LỤC MA TRẬN YÊU CẦU", bold=True, size=16).font = Font(name=FONT, bold=True, size=16, color=GREEN_D)
    _o(ws, 2, 1, f"{kh.get('ten','')} · {d.get('ma_kh','')} · {tl.get('phien_ban','')} · {tl.get('ngay','')}", size=11)
    _o(ws, 3, 1, "MẬT — chỉ lưu hành giữa hai bên", size=10).font = Font(name=FONT, size=10, color=RED, bold=True)
    ws.column_dimensions["A"].width = 34; ws.column_dimensions["B"].width = 78

    r = 5
    _o(ws, r, 1, "Quý vị cần làm gì với tệp này", bold=True, fill=GREEN_50); _o(ws, r, 2, "", fill=GREEN_50); r += 1
    for a, b in [
        ("Sheet 01 · Danh mục Yêu cầu", "Rà từng dòng. Ghi nhận xét vào cột “Ý kiến của Quý vị” và chọn ở cột “Kết luận”."),
        ("Sheet 02 · Chốt ưu tiên", "Bảng đếm tự động theo lựa chọn ở sheet 01 — không cần nhập tay."),
        ("Sheet 03 · Khoảng trống", "Những yêu cầu nền tảng chưa đáp ứng ngay và phương án đề xuất."),
        ("Sheet 04 · Kiến trúc Tác vụ", "Bổ sung danh mục công việc của từng phòng, mỗi tác vụ gắn một mã chỉ số."),
        ("Sheet 05 · Hệ nguồn dữ liệu", "Xác nhận hệ thống nào cung cấp số liệu và ai là chủ sở hữu."),
        ("Sheet 06 · Câu hỏi mở", "Trả lời trực tiếp vào cột dành riêng."),
    ]:
        _o(ws, r, 1, a, bold=True); _o(ws, r, 2, b); ws.row_dimensions[r].height = 30; r += 1

    r += 1
    _o(ws, r, 1, "Chú giải mức ưu tiên", bold=True, fill=GREEN_50); _o(ws, r, 2, "", fill=GREEN_50); r += 1
    for k, v in [("Bắt buộc", "Thiếu hạng mục này thì hệ thống không dùng được cho nghiệp vụ"),
                 ("Nên có", "Quan trọng, có thể lùi sang đợt sau mà nghiệp vụ vẫn chạy"),
                 ("Có thì tốt", "Làm nếu còn nguồn lực"),
                 ("Đợt sau", "Hai bên thống nhất chưa làm lần này")]:
        _o(ws, r, 1, k, bold=True); _o(ws, r, 2, v); r += 1

    r += 1
    _o(ws, r, 1, "Chú giải mức đáp ứng", bold=True, fill=GREEN_50); _o(ws, r, 2, "", fill=GREEN_50); r += 1
    for k, v in [("Sẵn sàng", "Dùng được ngay sau khi khởi tạo hệ thống"),
                 ("Sẵn sàng — cần kích hoạt", "Đã có sẵn, chờ Quý tổ chức cung cấp hạ tầng hoặc quyền truy cập"),
                 ("Đáp ứng bằng cấu hình", "Thiết lập theo đặc thù của Quý tổ chức, không sửa mã nguồn"),
                 ("Cần phát triển", "Cần bổ sung; nguồn lực nêu ở Đề xuất giải pháp"),
                 ("Ngoài phạm vi đợt này", "Không nằm trong cam kết lần này")]:
        _o(ws, r, 1, k, bold=True, fill=FIT_MAU[k]); _o(ws, r, 2, v); r += 1

    # ── 01 Danh mục yêu cầu ──
    S1 = "01 · Danh mục Yêu cầu"
    ws = wb.create_sheet(S1)
    cot = ["Mã", "Phân hệ", "Tiêu đề", "Phát biểu yêu cầu", "Tiêu chí chấp nhận",
           "Ưu tiên", "Mức đáp ứng", "Khoảng trống", "Ý kiến của Quý vị", "Kết luận"]
    _dau(ws, cot, [11, 20, 24, 52, 46, 12, 24, 34, 34, 17])
    for i, y in enumerate(yc, start=2):
        _o(ws, i, 1, y.get("ma", ""), bold=True)
        _o(ws, i, 2, f"{y.get('phan_he','')} · {PHAN_HE.get(y.get('phan_he',''),'')}")
        _o(ws, i, 3, y.get("tieu_de", ""))
        _o(ws, i, 4, y.get("phat_bieu", ""))
        _o(ws, i, 5, y.get("tieu_chi_chap_nhan", ""))
        _o(ws, i, 6, MOSCOW.get(y.get("moscow", ""), y.get("moscow", "")),
           fill=RED_50 if y.get("moscow") == "M" else None)
        _o(ws, i, 7, y.get("muc_fit", ""), fill=FIT_MAU.get(y.get("muc_fit", "")))
        _o(ws, i, 8, y.get("gap", ""))
        _o(ws, i, 9, "")
        _o(ws, i, 10, "")
        ws.row_dimensions[i].height = 58
    n = len(yc) + 1
    dv = DataValidation(type="list", formula1='"Đồng ý,Cần trao đổi,Không đồng ý,Bỏ khỏi phạm vi"', allow_blank=True)
    ws.add_data_validation(dv); dv.add(f"J2:J{max(n, 2)}")
    dvp = DataValidation(type="list", formula1='"Bắt buộc,Nên có,Có thì tốt,Đợt sau"', allow_blank=True)
    ws.add_data_validation(dvp); dvp.add(f"F2:F{max(n, 2)}")
    ws.freeze_panes = "C2"
    ws.auto_filter.ref = f"A1:J{max(n, 2)}"
    ws.sheet_view.showGridLines = False

    # ── 02 Chốt ưu tiên (công thức sống) ──
    ws = wb.create_sheet("02 · Chốt ưu tiên")
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 34
    for c in "BCDEF":
        ws.column_dimensions[c].width = 14
    _o(ws, 1, 1, "BẢNG CHỐT ƯU TIÊN", bold=True, size=14).font = Font(name=FONT, bold=True, size=14, color=GREEN_D)
    _o(ws, 2, 1, "Các ô đếm tự động theo sheet 01 — Quý vị sửa cột Ưu tiên hoặc Kết luận ở sheet 01 thì số ở đây tự đổi.",
       size=9).font = Font(name=FONT, size=9, italic=True, color=INK_3)
    rng = f"'{S1}'"

    r = 4
    _dau(ws, ["Phân hệ", "Bắt buộc", "Nên có", "Có thì tốt", "Đợt sau", "Tổng"], [34, 14, 14, 14, 14, 14], dong=r)
    ma_ph = sorted({y.get("phan_he", "") for y in yc if y.get("phan_he")})
    for i, ph in enumerate(ma_ph, start=r + 1):
        _o(ws, i, 1, f"{ph} · {PHAN_HE.get(ph,'')}", bold=True)
        for j, uu in enumerate(["Bắt buộc", "Nên có", "Có thì tốt", "Đợt sau"], start=2):
            _o(ws, i, j, f'=COUNTIFS({rng}!$B:$B,"{ph} · *",{rng}!$F:$F,"{uu}")', wrap=False)
        _o(ws, i, 6, f"=SUM(B{i}:E{i})", wrap=False, bold=True)
    tong = r + len(ma_ph) + 1
    _o(ws, tong, 1, "TỔNG CỘNG", bold=True, fill=GREEN_50)
    for j in range(2, 7):
        col = get_column_letter(j)
        _o(ws, tong, j, f"=SUM({col}{r+1}:{col}{tong-1})", wrap=False, bold=True, fill=GREEN_50)

    r = tong + 3
    _dau(ws, ["Mức đáp ứng", "Bắt buộc", "Nên có", "Có thì tốt", "Đợt sau", "Tổng"], [34, 14, 14, 14, 14, 14], dong=r)
    for i, fit in enumerate(FIT_MAU, start=r + 1):
        _o(ws, i, 1, fit, bold=True, fill=FIT_MAU[fit])
        for j, uu in enumerate(["Bắt buộc", "Nên có", "Có thì tốt", "Đợt sau"], start=2):
            _o(ws, i, j, f'=COUNTIFS({rng}!$G:$G,"{fit}",{rng}!$F:$F,"{uu}")', wrap=False)
        _o(ws, i, 6, f"=SUM(B{i}:E{i})", wrap=False, bold=True)
    r2 = r + len(FIT_MAU) + 2
    _o(ws, r2, 1, "Điểm cần chú ý: ô giao giữa “Bắt buộc” và “Cần phát triển” là phần quyết định nguồn lực của đợt triển khai.",
       size=9).font = Font(name=FONT, size=9, italic=True, color=INK_3)

    r3 = r2 + 2
    _dau(ws, ["Kết luận của Quý vị", "Số yêu cầu"], [34, 16], dong=r3)
    for i, kl in enumerate(["Đồng ý", "Cần trao đổi", "Không đồng ý", "Bỏ khỏi phạm vi"], start=r3 + 1):
        _o(ws, i, 1, kl, bold=True)
        _o(ws, i, 2, f'=COUNTIF({rng}!$J:$J,"{kl}")', wrap=False)

    # ── 03 Khoảng trống ──
    ws = wb.create_sheet("03 · Khoảng trống")
    ws.sheet_view.showGridLines = False
    _dau(ws, ["Mã", "Yêu cầu", "Mức đáp ứng", "Khoảng trống", "Phương án đề xuất", "Ý kiến của Quý vị"],
         [11, 34, 24, 44, 40, 32])
    gaps = [y for y in yc if y.get("gap")]
    for i, y in enumerate(gaps, start=2):
        _o(ws, i, 1, y.get("ma", ""), bold=True)
        _o(ws, i, 2, y.get("tieu_de", ""))
        _o(ws, i, 3, y.get("muc_fit", ""), fill=FIT_MAU.get(y.get("muc_fit", "")))
        _o(ws, i, 4, y.get("gap", ""))
        _o(ws, i, 5, "")
        _o(ws, i, 6, "")
        ws.row_dimensions[i].height = 46
    ws.freeze_panes = "A2"

    # ── 04 Kiến trúc Tác vụ (Từ điển Tác vụ) ──
    ws = wb.create_sheet("04 · Kiến trúc Tác vụ")
    ws.sheet_view.showGridLines = False
    ktv = d.get("kien_truc_tac_vu", {})
    _o(ws, 1, 1, "KIẾN TRÚC TÁC VỤ — đầu vào cho Từ điển Tác vụ", bold=True, size=13).font = \
        Font(name=FONT, bold=True, size=13, color=GREEN_D)
    _o(ws, 2, 1, ktv.get("ghi_chu", ""), size=9).font = Font(name=FONT, size=9, italic=True, color=INK_3)
    _o(ws, 3, 1, "Mỗi tác vụ phải gắn một mã chỉ số. Tác vụ chưa gắn chỉ số sẽ không được đưa vào sử dụng — "
                 "đây là ràng buộc có chủ đích để kết quả đánh giá luôn giải thích được bằng công việc có thật.",
       size=9).font = Font(name=FONT, size=9, italic=True, color=INK_3)
    _dau(ws, ["Phòng / Bộ phận", "Nhóm tác vụ", "Tên tác vụ", "Người chịu trách nhiệm",
              "Đầu ra", "Mã chỉ số (KPI)", "Tần suất", "Trạng thái"],
         [20, 24, 40, 24, 30, 16, 14, 16], dong=5)
    dong = ktv.get("dong", [])
    for i, x in enumerate(dong, start=6):
        for j, k in enumerate(["phong", "nhom_tac_vu", "ten_tac_vu", "nguoi_chiu_trach_nhiem",
                               "dau_ra", "ma_kpi", "tan_suat", "trang_thai"], start=1):
            _o(ws, i, j, x.get(k, ""))
    het = 5 + max(len(dong), 1)
    for i in range(het + 1, het + 41):       # 40 dòng trống để khách bổ sung
        for j in range(1, 9):
            _o(ws, i, j, "")
    dvt = DataValidation(type="list", formula1='"Đề xuất,Đang dùng,Cần sửa,Ngừng dùng"', allow_blank=True)
    ws.add_data_validation(dvt); dvt.add(f"H6:H{het+40}")
    ws.freeze_panes = "A6"
    ws.auto_filter.ref = f"A5:H{het+40}"

    # ── 05 Hệ nguồn ──
    ws = wb.create_sheet("05 · Hệ nguồn dữ liệu")
    ws.sheet_view.showGridLines = False
    _dau(ws, ["Hệ nguồn", "Chủ sở hữu (chức danh)", "Dữ liệu cần lấy", "Có kết nối?",
              "Tần suất mong muốn", "Phương án dự phòng", "Xác nhận của Quý vị"],
         [26, 26, 34, 16, 20, 26, 28])
    hn = d.get("he_nguon", [])
    for i, x in enumerate(hn, start=2):
        for j, k in enumerate(["he", "chu_so_huu", "du_lieu", "co_api", "tan_suat", "du_phong"], start=1):
            _o(ws, i, j, x.get(k, ""))
        _o(ws, i, 7, "")
    b = len(hn) + 2
    for i in range(b, b + 15):
        for j in range(1, 8):
            _o(ws, i, j, "")
    dva = DataValidation(type="list", formula1='"Có,Không,Chưa rõ"', allow_blank=True)
    ws.add_data_validation(dva); dva.add(f"D2:D{b+14}")
    ws.freeze_panes = "A2"

    # ── 06 Câu hỏi mở ──
    ws = wb.create_sheet("06 · Câu hỏi mở")
    ws.sheet_view.showGridLines = False
    _dau(ws, ["#", "Câu hỏi", "Người trả lời", "Hạn", "Chặn yêu cầu", "TRẢ LỜI CỦA QUÝ VỊ"],
         [6, 56, 24, 18, 16, 56])
    for i, x in enumerate(d.get("cau_hoi_mo", []), start=2):
        _o(ws, i, 1, x.get("stt", ""))
        _o(ws, i, 2, x.get("noi_dung", ""))
        _o(ws, i, 3, x.get("nguoi_tra_loi", ""))
        _o(ws, i, 4, x.get("han", ""))
        _o(ws, i, 5, x.get("chan_yeu_cau", ""))
        _o(ws, i, 6, "", fill="FFFDF5")
        ws.row_dimensions[i].height = 44
    ws.freeze_panes = "A2"

    ra.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(ra))


# ─────────────────────────── chạy ───────────────────────────
def main():
    try:                                   # console Windows mặc định cp1252 → vỡ tiếng Việt
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass
    ap = argparse.ArgumentParser(description="Sinh bộ hồ sơ BRD gửi khách từ ho-so.json")
    ap.add_argument("nguon", help="đường dẫn ho-so.json")
    ap.add_argument("--out", default=None, help="thư mục đầu ra (mặc định: cùng chỗ với nguồn)")
    a = ap.parse_args()

    src = Path(a.nguon).resolve()
    raw = json.loads(src.read_text(encoding="utf-8"))
    d = bo_khoa_noi_bo(raw)

    loi, canh = quet_hang_rao(d)
    if loi:
        print("⛔ HÀNG RÀO CHẶN — không ghi file. Dấu vết nội bộ còn sót:", file=sys.stderr)
        for duong, tu, trich in loi:
            print(f"   {duong}: khớp “{tu}” trong «{trich}…»", file=sys.stderr)
        sys.exit(2)

    out = Path(a.out).resolve() if a.out else src.parent
    ma, pb = d.get("ma_kh", "KH"), d.get("tai_lieu", {}).get("phien_ban", "v1")
    f_doc = out / f"01_BRD_{ma}_{pb}.docx"
    f_xls = out / f"02_Phu_luc_{ma}_{pb}.xlsx"

    dung_docx(d, f_doc)
    dung_xlsx(d, f_xls)

    print("✅ Hàng rào chặn cứng: SẠCH (0 dấu vết nội bộ)")
    if canh:
        print(f"⚠  {len(canh)} chỗ cần liếc mắt xác nhận là nội dung CỦA KHÁCH, không phải ước lượng của mình:")
        for duong, tu, trich in canh:
            print(f"   {duong}: “{tu}” trong «{trich}…»")
    print(f"   {f_doc.name}   ({f_doc.stat().st_size:,} bytes)")
    print(f"   {f_xls.name}   ({f_xls.stat().st_size:,} bytes)")
    print(f"   Thư mục: {out}")
    print(f"   Yêu cầu: {len(d.get('yeu_cau', []))} · Tác vụ: {len(d.get('kien_truc_tac_vu', {}).get('dong', []))}")


if __name__ == "__main__":
    main()
